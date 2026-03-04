import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'D:\SKRIPSI\Skripsi_Bevan\Draft Bepan (2).docx')
with open(r'D:\SKRIPSI\Skripsi_Bevan\_doc_extract.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            style = p.style.name if p.style else 'None'
            f.write(f'[{i}][{style}] {text}\n')
    # Also extract tables
    for ti, table in enumerate(doc.tables):
        f.write(f'\n=== TABLE {ti} ===\n')
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            separator = " | "
            f.write(f'  Row {ri}: {separator.join(cells)}\n')

print('Done. Lines:', sum(1 for _ in open(r'D:\SKRIPSI\Skripsi_Bevan\_doc_extract.txt', encoding='utf-8')))
